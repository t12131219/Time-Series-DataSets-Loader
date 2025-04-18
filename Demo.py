import pandas as pd
from docx import Document
from docx.shared import Pt

# 读取 Word 文件
word_file = '数据集总结.docx'
doc = Document(word_file)

# 提取 Word 文件中的数据
word_data = {}
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            name = cells[0].text.strip()
            description = cells[1].text.strip()
            word_data[name] = description

# 读取 Excel 文件
excel_file = '新建 Microsoft Excel 工作表.xlsx'
excel_df = pd.read_excel(excel_file)

# 提取 Excel 文件中的数据
excel_data = {}
for index, row in excel_df.iterrows():
    name = row['Name']
    description = row['Description']
    excel_data[name] = description

# 匹配和整合数据
merged_data = {}
for name in word_data:
    if name in excel_data:
        merged_description = f"{word_data[name]}\n\n{excel_data[name]}"
    else:
        merged_description = word_data[name]
    merged_data[name] = merged_description

# 生成新的 Word 文件
new_doc = Document()
new_doc.add_heading('数据集总结', 0)

for name in merged_data:
    new_doc.add_heading(name, level=1)
    new_paragraph = new_doc.add_paragraph(merged_data[name])
    for run in new_paragraph.runs:
        run.font.size = Pt(12)

new_doc.save('新的数据集总结.docx')
print("新的 Word 文件已生成！")